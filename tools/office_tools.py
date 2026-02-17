"""
Advanced Office Document Generation Tools
Supports professional Word documents with rich styling, charts, images, TOC, headers/footers,
and professional Excel documents with charts, conditional formatting, and advanced styling.
"""
import os
import io
import json
import tempfile
import math

# --- Word Document Imports ---
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# --- Excel Imports ---
import openpyxl
from openpyxl.styles import (
    Font, Alignment, PatternFill, Border, Side, numbers, NamedStyle
)
from openpyxl.utils import get_column_letter
from openpyxl.chart import (
    BarChart, LineChart, PieChart, AreaChart, ScatterChart,
    Reference, Series
)
from openpyxl.chart.label import DataLabelList
from openpyxl.formatting.rule import (
    CellIsRule, ColorScaleRule, DataBarRule, IconSetRule
)
from openpyxl.worksheet.datavalidation import DataValidation

# --- Chart Generation (matplotlib) ---
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np

from .base import get_workspace_path

# ============================================================================
#  UTILITY: Color & Style Helpers
# ============================================================================

def _hex_to_rgb(hex_color):
    """Convert '#RRGGBB' or 'RRGGBB' to RGBColor."""
    hex_color = hex_color.lstrip('#')
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

def _parse_color(color_str, default="333333"):
    """Parse color string to hex without '#'. Accepts '#RRGGBB', 'RRGGBB', or named colors."""
    if not color_str:
        return default
    color_str = color_str.strip().lstrip('#')
    named = {
        "red": "FF0000", "green": "00AA00", "blue": "0066CC", "black": "000000",
        "white": "FFFFFF", "gray": "888888", "grey": "888888", "orange": "FF8800",
        "purple": "8800CC", "yellow": "FFCC00", "navy": "003366", "teal": "008080",
        "darkblue": "003366", "darkgreen": "006600", "darkred": "990000",
    }
    if color_str.lower() in named:
        return named[color_str.lower()]
    if len(color_str) == 6:
        return color_str
    return default

# ============================================================================
#  UTILITY: Generate chart image via matplotlib
# ============================================================================

def _generate_chart_image(chart_config, workspace):
    """
    Generate a chart image using matplotlib and return the file path.
    chart_config keys:
      - chart_type: bar, line, pie, scatter, area, bar_horizontal, grouped_bar, stacked_bar
      - title: chart title
      - x_label, y_label: axis labels
      - categories: list of category labels (x-axis)
      - series: list of {"name": str, "values": list[number]}
      - colors: optional list of hex colors
      - width, height: figure size in inches (default 8x5)
      - style: optional matplotlib style name
      - show_legend: bool (default True)
      - show_values: bool (default False) — show value labels on bars/points
      - font_size: base font size (default 11)
    """
    chart_type = chart_config.get("chart_type", "bar")
    title = chart_config.get("title", "")
    x_label = chart_config.get("x_label", "")
    y_label = chart_config.get("y_label", "")
    categories = chart_config.get("categories", [])
    series_list = chart_config.get("series", [])
    colors = chart_config.get("colors", None)
    width = chart_config.get("width", 8)
    height = chart_config.get("height", 5)
    style = chart_config.get("style", "seaborn-v0_8-whitegrid")
    show_legend = chart_config.get("show_legend", True)
    show_values = chart_config.get("show_values", False)
    font_size = chart_config.get("font_size", 11)

    # Default professional color palette
    default_colors = [
        "#4472C4", "#ED7D31", "#A5A5A5", "#FFC000", "#5B9BD5",
        "#70AD47", "#264478", "#9B57A0", "#636363", "#EB7E30"
    ]
    if not colors:
        colors = default_colors

    try:
        plt.style.use(style)
    except:
        pass

    plt.rcParams.update({
        'font.size': font_size,
        'axes.titlesize': font_size + 3,
        'axes.labelsize': font_size + 1,
        'figure.dpi': 150,
    })

    # Try to use a CJK font if available
    # Added 'WenQuanYi Zen Hei' and 'STHeiti' for better compatibility
    cjk_fonts = ['SimHei', 'Microsoft YaHei', 'WenQuanYi Micro Hei', 'WenQuanYi Zen Hei', 'Noto Sans CJK SC', 'STHeiti', 'DejaVu Sans']
    
    # Force rebuild font cache to ensure new fonts are detected
    try:
        fm._get_fontconfig_fonts.cache_clear()
    except:
        pass

    for font_name in cjk_fonts:
        try:
            # Check if font is available in matplotlib's font manager
            # We use a more robust check by looking at the names in the font manager
            if any(font_name.lower() in f.name.lower() for f in fm.fontManager.ttflist):
                plt.rcParams['font.sans-serif'] = [font_name] + plt.rcParams.get('font.sans-serif', [])
                break
            # Fallback to findfont check
            fm.findfont(font_name, fallback_to_default=False)
            plt.rcParams['font.sans-serif'] = [font_name] + plt.rcParams.get('font.sans-serif', [])
            break
        except:
            continue
    plt.rcParams['axes.unicode_minus'] = False

    fig, ax = plt.subplots(figsize=(width, height))

    if chart_type == "pie":
        if series_list:
            values = series_list[0].get("values", [])
            explode = [0.03] * len(values)
            wedges, texts, autotexts = ax.pie(
                values, labels=categories, autopct='%1.1f%%',
                colors=[c for c in colors[:len(values)]],
                explode=explode, shadow=False, startangle=90,
                textprops={'fontsize': font_size}
            )
            for t in autotexts:
                t.set_fontsize(font_size - 1)
        ax.set_title(title, fontsize=font_size + 3, fontweight='bold', pad=15)

    elif chart_type == "scatter":
        for i, s in enumerate(series_list):
            vals = s.get("values", [])
            x_vals = list(range(len(vals))) if not categories else list(range(len(categories)))
            ax.scatter(x_vals[:len(vals)], vals, label=s.get("name", f"Series {i+1}"),
                      color=colors[i % len(colors)], s=60, alpha=0.8, edgecolors='white', linewidth=0.5)
        if categories:
            ax.set_xticks(range(len(categories)))
            ax.set_xticklabels(categories, rotation=45 if len(categories) > 6 else 0, ha='right')

    elif chart_type == "area":
        x = list(range(len(categories))) if categories else []
        for i, s in enumerate(series_list):
            vals = s.get("values", [])
            ax.fill_between(x[:len(vals)], vals, alpha=0.4, color=colors[i % len(colors)],
                           label=s.get("name", f"Series {i+1}"))
            ax.plot(x[:len(vals)], vals, color=colors[i % len(colors)], linewidth=1.5)
        if categories:
            ax.set_xticks(range(len(categories)))
            ax.set_xticklabels(categories, rotation=45 if len(categories) > 6 else 0, ha='right')

    elif chart_type == "bar_horizontal":
        y_pos = np.arange(len(categories))
        bar_height = 0.8 / max(len(series_list), 1)
        for i, s in enumerate(series_list):
            vals = s.get("values", [])
            offset = (i - len(series_list)/2 + 0.5) * bar_height
            bars = ax.barh(y_pos + offset, vals[:len(categories)], bar_height,
                          label=s.get("name", f"Series {i+1}"),
                          color=colors[i % len(colors)], edgecolor='white', linewidth=0.5)
            if show_values:
                for bar in bars:
                    w = bar.get_width()
                    ax.text(w + max(vals)*0.01, bar.get_y() + bar.get_height()/2,
                           f'{w:,.0f}', va='center', fontsize=font_size - 2)
        ax.set_yticks(y_pos)
        ax.set_yticklabels(categories)

    elif chart_type in ("stacked_bar",):
        x = np.arange(len(categories))
        bar_width = 0.6
        bottom = np.zeros(len(categories))
        for i, s in enumerate(series_list):
            vals = np.array(s.get("values", [0]*len(categories))[:len(categories)], dtype=float)
            ax.bar(x, vals, bar_width, bottom=bottom,
                  label=s.get("name", f"Series {i+1}"),
                  color=colors[i % len(colors)], edgecolor='white', linewidth=0.5)
            bottom += vals
        ax.set_xticks(x)
        ax.set_xticklabels(categories, rotation=45 if len(categories) > 6 else 0, ha='right')

    elif chart_type == "line":
        x = list(range(len(categories))) if categories else []
        for i, s in enumerate(series_list):
            vals = s.get("values", [])
            ax.plot(x[:len(vals)], vals, marker='o', markersize=5,
                   label=s.get("name", f"Series {i+1}"),
                   color=colors[i % len(colors)], linewidth=2)
            if show_values:
                for xi, vi in zip(x[:len(vals)], vals):
                    ax.annotate(f'{vi:,.0f}', (xi, vi), textcoords="offset points",
                               xytext=(0, 8), ha='center', fontsize=font_size - 2)
        if categories:
            ax.set_xticks(range(len(categories)))
            ax.set_xticklabels(categories, rotation=45 if len(categories) > 6 else 0, ha='right')

    else:  # bar (default), grouped_bar
        x = np.arange(len(categories))
        n_series = max(len(series_list), 1)
        bar_width = 0.8 / n_series
        for i, s in enumerate(series_list):
            vals = s.get("values", [])
            offset = (i - n_series/2 + 0.5) * bar_width
            bars = ax.bar(x + offset, vals[:len(categories)], bar_width,
                         label=s.get("name", f"Series {i+1}"),
                         color=colors[i % len(colors)], edgecolor='white', linewidth=0.5)
            if show_values:
                for bar in bars:
                    h = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2, h + max(vals)*0.01,
                           f'{h:,.0f}', ha='center', va='bottom', fontsize=font_size - 2)
        ax.set_xticks(x)
        ax.set_xticklabels(categories, rotation=45 if len(categories) > 6 else 0, ha='right')

    if chart_type != "pie":
        if title:
            ax.set_title(title, fontsize=font_size + 3, fontweight='bold', pad=12)
        if x_label:
            ax.set_xlabel(x_label)
        if y_label:
            ax.set_ylabel(y_label)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    if show_legend and len(series_list) > 1 and chart_type != "pie":
        ax.legend(frameon=True, fancybox=True, shadow=False, fontsize=font_size - 1)

    plt.tight_layout()

    # Save to workspace temp directory
    charts_dir = os.path.join(workspace, ".charts")
    os.makedirs(charts_dir, exist_ok=True)
    import uuid
    chart_path = os.path.join(charts_dir, f"chart_{uuid.uuid4().hex[:8]}.png")
    fig.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(fig)
    return chart_path


# ============================================================================
#  WORD DOCUMENT GENERATION
# ============================================================================

def _apply_paragraph_format(paragraph, fmt):
    """Apply formatting to a paragraph from a format dict."""
    if not fmt:
        return
    pf = paragraph.paragraph_format
    if "alignment" in fmt:
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        pf.alignment = align_map.get(fmt["alignment"], WD_ALIGN_PARAGRAPH.LEFT)
    if "space_before" in fmt:
        pf.space_before = Pt(fmt["space_before"])
    if "space_after" in fmt:
        pf.space_after = Pt(fmt["space_after"])
    if "line_spacing" in fmt:
        pf.line_spacing = fmt["line_spacing"]
    if "first_line_indent" in fmt:
        pf.first_line_indent = Cm(fmt["first_line_indent"])
    if "left_indent" in fmt:
        pf.left_indent = Cm(fmt["left_indent"])


def _apply_run_format(run, fmt):
    """Apply formatting to a run from a format dict."""
    if not fmt:
        return
    if "bold" in fmt:
        run.bold = fmt["bold"]
    if "italic" in fmt:
        run.italic = fmt["italic"]
    if "underline" in fmt:
        run.underline = fmt["underline"]
    if "strike" in fmt:
        run.font.strike = fmt["strike"]
    if "font_size" in fmt:
        run.font.size = Pt(fmt["font_size"])
    if "font_name" in fmt:
        run.font.name = fmt["font_name"]
        # Also set East Asian font
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{fmt["font_name"]}"/>')
            rPr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), fmt["font_name"])
    if "color" in fmt:
        run.font.color.rgb = _hex_to_rgb(fmt["color"])
    if "highlight" in fmt:
        run.font.highlight_color = fmt["highlight"]
    if "superscript" in fmt:
        run.font.superscript = fmt["superscript"]
    if "subscript" in fmt:
        run.font.subscript = fmt["subscript"]


def _add_rich_text(paragraph, text_parts):
    """
    Add rich text to a paragraph.
    text_parts can be:
      - a string (plain text)
      - a list of {"text": str, ...formatting options}
    """
    if isinstance(text_parts, str):
        run = paragraph.add_run(text_parts)
        return
    for part in text_parts:
        if isinstance(part, str):
            paragraph.add_run(part)
        elif isinstance(part, dict):
            run = paragraph.add_run(part.get("text", ""))
            _apply_run_format(run, part)


def _style_table_cell(cell, bg_color=None, font_color=None, bold=False, font_size=None, alignment="center"):
    """Style a single table cell."""
    if bg_color:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{_parse_color(bg_color)}"/>')
        cell._element.get_or_add_tcPr().append(shading)
    for paragraph in cell.paragraphs:
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        paragraph.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
        for run in paragraph.runs:
            if font_color:
                run.font.color.rgb = _hex_to_rgb(font_color)
            if bold:
                run.bold = True
            if font_size:
                run.font.size = Pt(font_size)


def _add_table_of_contents(doc):
    """Add a Table of Contents field to the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar2)

    run4 = paragraph.add_run("(Right-click and select 'Update Field' to refresh TOC)")
    run4.italic = True
    run4.font.color.rgb = RGBColor(128, 128, 128)

    run5 = paragraph.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar3)


def _set_header_footer(section, header_text=None, footer_text=None, header_style=None, footer_style=None):
    """Set header and footer for a section."""
    if header_text:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = header_text
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
        if header_style:
            _apply_run_format(hp.runs[0] if hp.runs else hp.add_run(header_text), header_style)

    if footer_text:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = footer_text
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)


def _add_page_number(section):
    """Add page number to footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)

    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)

    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def generate_word_document(file_path, content_structure, document_settings=None):
    """
    Generate a professional Word document with advanced features.

    Args:
        file_path: Path to save the .docx file (relative to workspace)
        content_structure: list of content blocks, each is a dict with "type" and type-specific keys.
        document_settings: optional dict for document-level settings.

    Supported content block types:
        - heading: {type, text, level(1-9), format{alignment, color, font_name, font_size}}
        - paragraph: {type, text(str or list of rich parts), format{alignment, line_spacing, first_line_indent, ...}}
        - rich_paragraph: {type, parts:[{text, bold, italic, color, font_size, font_name, underline, ...}], format{...}}
        - table: {type, rows(2D array), header_bg_color, header_font_color, stripe_colors[even,odd],
                  col_widths[inches], style, border_color, font_size, alignment}
        - image: {type, path(relative to workspace), width(inches), alignment}
        - chart: {type:"chart", chart_type, title, categories, series, colors, width, height, ...}
        - page_break: {type}
        - toc: {type} — insert Table of Contents
        - horizontal_rule: {type}
        - bullet_list: {type, items:[str or {text, level(0-2)}]}
        - numbered_list: {type, items:[str or {text, level(0-2)}]}
        - quote: {type, text, author}
        - code_block: {type, code, language}
        - cover_page: {type, title, subtitle, author, date, logo_path, bg_color}
        - section_break: {type, orientation("portrait"/"landscape")}
        - columns: {type, count(2-3)} — set columns for current section
        - footnote: {type, text, reference_text}
        - watermark: {type, text}

    document_settings keys:
        - default_font: str
        - default_font_size: int (pt)
        - margins: {top, bottom, left, right} in cm
        - header_text: str
        - footer_text: str
        - page_numbers: bool
        - line_spacing: float
        - orientation: "portrait" or "landscape"
    """
    workspace = get_workspace_path()
    if not workspace:
        return {"error": "Workspace not configured"}

    full_path = os.path.join(workspace, file_path)
    os.makedirs(os.path.dirname(full_path) if os.path.dirname(full_path) else workspace, exist_ok=True)

    doc = Document()
    settings = document_settings or {}

    # --- Apply document-level settings ---
    style = doc.styles['Normal']
    font = style.font
    font.name = settings.get("default_font", "Calibri")
    font.size = Pt(settings.get("default_font_size", 11))
    # Set East Asian font
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font.name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font.name)

    if "line_spacing" in settings:
        style.paragraph_format.line_spacing = settings["line_spacing"]

    # Page margins
    section = doc.sections[0]
    margins = settings.get("margins", {})
    if "top" in margins: section.top_margin = Cm(margins["top"])
    if "bottom" in margins: section.bottom_margin = Cm(margins["bottom"])
    if "left" in margins: section.left_margin = Cm(margins["left"])
    if "right" in margins: section.right_margin = Cm(margins["right"])

    if settings.get("orientation") == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

    # Header / Footer / Page numbers
    if settings.get("header_text") or settings.get("footer_text"):
        _set_header_footer(section, settings.get("header_text"), settings.get("footer_text"))
    if settings.get("page_numbers"):
        _add_page_number(section)

    # --- Process content blocks ---
    try:
        for item in content_structure:
            item_type = item.get("type", "paragraph")
            fmt = item.get("format", {})

            if item_type == "cover_page":
                # Create a professional cover page
                for _ in range(4):
                    doc.add_paragraph()
                if item.get("logo_path"):
                    logo_full = os.path.join(workspace, item["logo_path"])
                    if os.path.exists(logo_full):
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(logo_full, width=Inches(2))
                        doc.add_paragraph()

                title_p = doc.add_paragraph()
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_p.add_run(item.get("title", ""))
                title_run.bold = True
                title_run.font.size = Pt(item.get("title_font_size", 36))
                title_run.font.color.rgb = _hex_to_rgb(item.get("title_color", "#003366"))

                if item.get("subtitle"):
                    sub_p = doc.add_paragraph()
                    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sub_p.paragraph_format.space_before = Pt(12)
                    sub_run = sub_p.add_run(item["subtitle"])
                    sub_run.font.size = Pt(item.get("subtitle_font_size", 18))
                    sub_run.font.color.rgb = _hex_to_rgb(item.get("subtitle_color", "#666666"))

                doc.add_paragraph()
                doc.add_paragraph()

                if item.get("author") or item.get("date"):
                    info_p = doc.add_paragraph()
                    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    info_text = ""
                    if item.get("author"):
                        info_text += item["author"]
                    if item.get("date"):
                        if info_text:
                            info_text += "\n"
                        info_text += item["date"]
                    info_run = info_p.add_run(info_text)
                    info_run.font.size = Pt(14)
                    info_run.font.color.rgb = RGBColor(100, 100, 100)

                doc.add_page_break()

            elif item_type == "toc":
                _add_table_of_contents(doc)
                doc.add_page_break()

            elif item_type == "heading":
                level = item.get("level", 1)
                h = doc.add_heading(item.get("text", ""), level=level)
                if fmt.get("color"):
                    for run in h.runs:
                        run.font.color.rgb = _hex_to_rgb(fmt["color"])
                if fmt.get("font_name"):
                    for run in h.runs:
                        run.font.name = fmt["font_name"]
                _apply_paragraph_format(h, fmt)

            elif item_type == "paragraph":
                p = doc.add_paragraph()
                text = item.get("text", "")
                if isinstance(text, list):
                    _add_rich_text(p, text)
                else:
                    run = p.add_run(text)
                    _apply_run_format(run, fmt)
                _apply_paragraph_format(p, fmt)

            elif item_type == "rich_paragraph":
                p = doc.add_paragraph()
                parts = item.get("parts", [])
                _add_rich_text(p, parts)
                _apply_paragraph_format(p, fmt)

            elif item_type == "bullet_list":
                for li in item.get("items", []):
                    if isinstance(li, str):
                        doc.add_paragraph(li, style='List Bullet')
                    elif isinstance(li, dict):
                        level = li.get("level", 0)
                        style_name = 'List Bullet' + (f' {level+1}' if level > 0 else '')
                        try:
                            p = doc.add_paragraph(li.get("text", ""), style=style_name)
                        except:
                            p = doc.add_paragraph(li.get("text", ""), style='List Bullet')
                            if level > 0:
                                p.paragraph_format.left_indent = Cm(level * 1.27)

            elif item_type == "numbered_list":
                for li in item.get("items", []):
                    if isinstance(li, str):
                        doc.add_paragraph(li, style='List Number')
                    elif isinstance(li, dict):
                        level = li.get("level", 0)
                        style_name = 'List Number' + (f' {level+1}' if level > 0 else '')
                        try:
                            p = doc.add_paragraph(li.get("text", ""), style=style_name)
                        except:
                            p = doc.add_paragraph(li.get("text", ""), style='List Number')
                            if level > 0:
                                p.paragraph_format.left_indent = Cm(level * 1.27)

            elif item_type == "table":
                rows_data = item.get("rows", [])
                if not rows_data:
                    continue
                n_rows = len(rows_data)
                n_cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=n_rows, cols=n_cols)

                # Apply table style
                tbl_style = item.get("style", "Table Grid")
                try:
                    table.style = tbl_style
                except:
                    table.style = "Table Grid"

                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Set column widths
                col_widths = item.get("col_widths", None)
                if col_widths:
                    for i, w in enumerate(col_widths):
                        if i < n_cols:
                            for row in table.rows:
                                row.cells[i].width = Inches(w)

                header_bg = item.get("header_bg_color", "#4472C4")
                header_fg = item.get("header_font_color", "#FFFFFF")
                stripe_colors = item.get("stripe_colors", ["#FFFFFF", "#F2F2F2"])
                cell_font_size = item.get("font_size", 10)
                cell_alignment = item.get("alignment", "center")

                for r_idx, row_data in enumerate(rows_data):
                    for c_idx in range(n_cols):
                        cell_val = row_data[c_idx] if c_idx < len(row_data) else ""
                        cell = table.cell(r_idx, c_idx)
                        cell.text = str(cell_val)

                        if r_idx == 0:
                            _style_table_cell(cell, bg_color=header_bg, font_color=header_fg,
                                            bold=True, font_size=cell_font_size, alignment=cell_alignment)
                        else:
                            bg = stripe_colors[r_idx % len(stripe_colors)] if stripe_colors else None
                            _style_table_cell(cell, bg_color=bg, font_size=cell_font_size,
                                            alignment=cell_alignment)

            elif item_type == "image":
                img_path = item.get("path", "")
                full_img = os.path.join(workspace, img_path)
                if os.path.exists(full_img):
                    p = doc.add_paragraph()
                    align = item.get("alignment", "center")
                    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}
                    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)
                    run = p.add_run()
                    run.add_picture(full_img, width=Inches(item.get("width", 5)))
                    # Caption
                    if item.get("caption"):
                        cap_p = doc.add_paragraph()
                        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_run = cap_p.add_run(item["caption"])
                        cap_run.italic = True
                        cap_run.font.size = Pt(9)
                        cap_run.font.color.rgb = RGBColor(100, 100, 100)

            elif item_type == "chart":
                chart_path = _generate_chart_image(item, workspace)
                if chart_path and os.path.exists(chart_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    chart_width = item.get("doc_width", 6)
                    run.add_picture(chart_path, width=Inches(chart_width))
                    if item.get("title"):
                        cap_p = doc.add_paragraph()
                        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_run = cap_p.add_run(item["title"])
                        cap_run.italic = True
                        cap_run.font.size = Pt(9)
                        cap_run.font.color.rgb = RGBColor(100, 100, 100)

            elif item_type == "quote":
                # Styled blockquote
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                # Add left border via XML
                pPr = p._element.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="4472C4"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)
                run = p.add_run(item.get("text", ""))
                run.italic = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(80, 80, 80)
                if item.get("author"):
                    author_p = doc.add_paragraph()
                    author_p.paragraph_format.left_indent = Cm(1.5)
                    author_run = author_p.add_run(f"— {item['author']}")
                    author_run.font.size = Pt(10)
                    author_run.font.color.rgb = RGBColor(120, 120, 120)

            elif item_type == "code_block":
                code = item.get("code", "")
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                # Gray background
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                p._element.get_or_add_pPr().append(shading)
                run = p.add_run(code)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(50, 50, 50)
                # Set East Asian font too
                r_elem = run._element
                rPr_code = r_elem.find(qn('w:rPr'))
                if rPr_code is None:
                    rPr_code = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                    r_elem.insert(0, rPr_code)
                rFonts_code = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Consolas" w:ascii="Consolas" w:hAnsi="Consolas"/>')
                rPr_code.insert(0, rFonts_code)

            elif item_type == "horizontal_rule":
                p = doc.add_paragraph()
                pPr = p._element.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)

            elif item_type == "page_break":
                doc.add_page_break()

            elif item_type == "section_break":
                new_section = doc.add_section()
                if item.get("orientation") == "landscape":
                    new_section.orientation = WD_ORIENT.LANDSCAPE
                    new_width, new_height = new_section.page_height, new_section.page_width
                    new_section.page_width = new_width
                    new_section.page_height = new_height
                # Carry over header/footer settings
                if settings.get("header_text") or settings.get("footer_text"):
                    _set_header_footer(new_section, settings.get("header_text"), settings.get("footer_text"))
                if settings.get("page_numbers"):
                    _add_page_number(new_section)

            elif item_type == "watermark":
                # Simple text watermark via header
                for sec in doc.sections:
                    header = sec.header
                    header.is_linked_to_previous = False
                    wp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                    wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    wr = wp.add_run(item.get("text", "DRAFT"))
                    wr.font.size = Pt(48)
                    wr.font.color.rgb = RGBColor(220, 220, 220)
                    wr.bold = True

        doc.save(full_path)
        return {"success": True, "file_path": full_path, "message": f"Word document saved to {file_path}"}
    except Exception as e:
        return {"error": f"Word generation error: {str(e)}"}


# ============================================================================
#  EXCEL DOCUMENT GENERATION
# ============================================================================

def _apply_cell_style(cell, style_dict):
    """Apply comprehensive styling to an Excel cell."""
    if not style_dict:
        return

    if "font_name" in style_dict or "font_size" in style_dict or "bold" in style_dict or "font_color" in style_dict or "italic" in style_dict:
        cell.font = Font(
            name=style_dict.get("font_name", cell.font.name),
            size=style_dict.get("font_size", cell.font.size),
            bold=style_dict.get("bold", cell.font.bold),
            italic=style_dict.get("italic", cell.font.italic),
            color=_parse_color(style_dict.get("font_color")) if style_dict.get("font_color") else cell.font.color,
            underline=style_dict.get("underline", cell.font.underline),
        )

    if "bg_color" in style_dict:
        cell.fill = PatternFill(
            start_color=_parse_color(style_dict["bg_color"]),
            end_color=_parse_color(style_dict["bg_color"]),
            fill_type="solid"
        )

    if "alignment" in style_dict or "wrap_text" in style_dict or "vertical" in style_dict:
        h_align = style_dict.get("alignment", "center")
        v_align = style_dict.get("vertical", "center")
        cell.alignment = Alignment(
            horizontal=h_align,
            vertical=v_align,
            wrap_text=style_dict.get("wrap_text", False),
            text_rotation=style_dict.get("text_rotation", 0)
        )

    if "number_format" in style_dict:
        cell.number_format = style_dict["number_format"]

    if "border" in style_dict:
        border_color = _parse_color(style_dict["border"].get("color", "000000"))
        border_style = style_dict["border"].get("style", "thin")
        side = Side(style=border_style, color=border_color)
        cell.border = Border(
            left=side if style_dict["border"].get("left", True) else Side(),
            right=side if style_dict["border"].get("right", True) else Side(),
            top=side if style_dict["border"].get("top", True) else Side(),
            bottom=side if style_dict["border"].get("bottom", True) else Side(),
        )


def generate_excel_document(file_path, sheets_data, workbook_settings=None):
    """
    Generate a professional Excel document with advanced features.

    Args:
        file_path: Path to save the .xlsx file (relative to workspace)
        sheets_data: list of sheet definitions
        workbook_settings: optional dict for workbook-level settings

    Each sheet definition:
    {
        "name": "Sheet Name",
        "data": [[row1], [row2], ...],
        "formulas": [{"cell": "A1", "formula": "=SUM(...)"}],
        "column_widths": {"A": 15, "B": 20, ...} or [15, 20, ...],
        "row_heights": {"1": 30, "2": 25, ...},
        "merge_cells": ["A1:C1", "D5:F5"],
        "header_style": {font_name, font_size, bold, font_color, bg_color, alignment, border},
        "data_style": {font_name, font_size, font_color, bg_color, alignment, border, wrap_text},
        "cell_styles": [{"range": "A1:C1", ...style_dict}, {"cell": "D5", ...style_dict}],
        "conditional_formatting": [
            {"range": "B2:B100", "type": "color_scale", "min_color": "FF0000", "max_color": "00FF00"},
            {"range": "C2:C100", "type": "data_bar", "color": "4472C4"},
            {"range": "D2:D100", "type": "cell_is", "operator": "greaterThan", "value": 100,
             "font_color": "00AA00", "bg_color": "CCFFCC"},
            {"range": "E2:E100", "type": "icon_set", "icon_style": "3Arrows"}
        ],
        "charts": [
            {
                "type": "bar",  // bar, line, pie, area, scatter
                "title": "Chart Title",
                "data_range": {"min_col": 1, "max_col": 3, "min_row": 1, "max_row": 10},
                "categories_range": {"min_col": 1, "min_row": 2, "max_row": 10},
                "position": "E2",
                "width": 15, "height": 10,
                "style": 10,
                "x_axis_title": "", "y_axis_title": ""
            }
        ],
        "data_validations": [
            {"range": "A2:A100", "type": "list", "formula": "Option1,Option2,Option3"},
            {"range": "B2:B100", "type": "whole", "min": 0, "max": 100}
        ],
        "freeze_panes": "A2",
        "auto_filter": "A1:F1",
        "print_settings": {"orientation": "landscape", "fit_to_page": true}
    }
    """
    workspace = get_workspace_path()
    if not workspace:
        return {"error": "Workspace not configured"}

    full_path = os.path.join(workspace, file_path)
    os.makedirs(os.path.dirname(full_path) if os.path.dirname(full_path) else workspace, exist_ok=True)

    wb = openpyxl.Workbook()
    wb_settings = workbook_settings or {}

    # Remove default sheet
    default_sheet = wb.active
    wb.remove(default_sheet)

    try:
        for sheet_info in sheets_data:
            ws = wb.create_sheet(title=sheet_info.get("name", "Sheet"))
            data = sheet_info.get("data", [])
            header_style = sheet_info.get("header_style", {
                "bold": True, "font_size": 11, "font_color": "FFFFFF",
                "bg_color": "4472C4", "alignment": "center",
                "border": {"color": "4472C4", "style": "thin"}
            })
            data_style = sheet_info.get("data_style", {})

            # --- Write data ---
            for r_idx, row in enumerate(data, 1):
                for c_idx, value in enumerate(row, 1):
                    cell = ws.cell(row=r_idx, column=c_idx, value=value)

                    if r_idx == 1 and header_style:
                        _apply_cell_style(cell, header_style)
                    elif data_style:
                        style_copy = data_style.copy()
                        # Alternate row coloring
                        if "stripe_colors" in sheet_info:
                            sc = sheet_info["stripe_colors"]
                            style_copy["bg_color"] = sc[0] if r_idx % 2 == 0 else sc[1]
                        _apply_cell_style(cell, style_copy)
                    elif r_idx == 1:
                        # Default header styling
                        _apply_cell_style(cell, {
                            "bold": True, "font_size": 11, "font_color": "FFFFFF",
                            "bg_color": "4472C4", "alignment": "center",
                            "border": {"color": "4472C4", "style": "thin"}
                        })
                    else:
                        # Default data styling: thin borders, center aligned
                        thin_side = Side(style='thin', color='D9D9D9')
                        cell.border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
                        cell.alignment = Alignment(horizontal='center', vertical='center')

            # --- Formulas ---
            for f in sheet_info.get("formulas", []):
                cell_ref = f.get("cell")
                formula = f.get("formula")
                if cell_ref and formula:
                    ws[cell_ref] = formula

            # --- Column widths ---
            col_widths = sheet_info.get("column_widths", None)
            if col_widths:
                if isinstance(col_widths, dict):
                    for col_letter, width in col_widths.items():
                        ws.column_dimensions[col_letter].width = width
                elif isinstance(col_widths, list):
                    for i, width in enumerate(col_widths):
                        ws.column_dimensions[get_column_letter(i + 1)].width = width
            else:
                # Auto-width based on content
                for c_idx in range(1, (max(len(r) for r in data) if data else 0) + 1):
                    max_len = 0
                    for r_idx in range(1, len(data) + 1):
                        cell = ws.cell(row=r_idx, column=c_idx)
                        if cell.value:
                            cell_len = len(str(cell.value))
                            if cell_len > max_len:
                                max_len = cell_len
                    ws.column_dimensions[get_column_letter(c_idx)].width = min(max(max_len + 4, 10), 50)

            # --- Row heights ---
            for row_num, height in sheet_info.get("row_heights", {}).items():
                ws.row_dimensions[int(row_num)].height = height

            # --- Merge cells ---
            for merge_range in sheet_info.get("merge_cells", []):
                ws.merge_cells(merge_range)

            # --- Cell-specific styles ---
            for cs in sheet_info.get("cell_styles", []):
                if "range" in cs:
                    style_dict = {k: v for k, v in cs.items() if k != "range"}
                    for row in ws[cs["range"]]:
                        for cell in (row if hasattr(row, '__iter__') else [row]):
                            _apply_cell_style(cell, style_dict)
                elif "cell" in cs:
                    style_dict = {k: v for k, v in cs.items() if k != "cell"}
                    _apply_cell_style(ws[cs["cell"]], style_dict)

            # --- Conditional Formatting ---
            for cf in sheet_info.get("conditional_formatting", []):
                cf_range = cf.get("range", "A1:A1")
                cf_type = cf.get("type", "")

                if cf_type == "color_scale":
                    rule = ColorScaleRule(
                        start_type='min', start_color=_parse_color(cf.get("min_color", "FF0000")),
                        end_type='max', end_color=_parse_color(cf.get("max_color", "00FF00"))
                    )
                    ws.conditional_formatting.add(cf_range, rule)

                elif cf_type == "data_bar":
                    rule = DataBarRule(
                        start_type='min', end_type='max',
                        color=_parse_color(cf.get("color", "4472C4"))
                    )
                    ws.conditional_formatting.add(cf_range, rule)

                elif cf_type == "cell_is":
                    rule = CellIsRule(
                        operator=cf.get("operator", "greaterThan"),
                        formula=[str(cf.get("value", 0))],
                        font=Font(color=_parse_color(cf.get("font_color", "000000"))) if cf.get("font_color") else None,
                        fill=PatternFill(
                            start_color=_parse_color(cf.get("bg_color", "FFFFFF")),
                            end_color=_parse_color(cf.get("bg_color", "FFFFFF")),
                            fill_type="solid"
                        ) if cf.get("bg_color") else None,
                    )
                    ws.conditional_formatting.add(cf_range, rule)

                elif cf_type == "icon_set":
                    rule = IconSetRule(
                        icon_style=cf.get("icon_style", "3Arrows"),
                        type='percent',
                        values=[0, 33, 67]
                    )
                    ws.conditional_formatting.add(cf_range, rule)

            # --- Charts ---
            for chart_def in sheet_info.get("charts", []):
                chart_type = chart_def.get("type", "bar")
                dr = chart_def.get("data_range", {})
                cr = chart_def.get("categories_range", {})

                if chart_type == "bar":
                    chart = BarChart()
                elif chart_type == "line":
                    chart = LineChart()
                elif chart_type == "pie":
                    chart = PieChart()
                elif chart_type == "area":
                    chart = AreaChart()
                elif chart_type == "scatter":
                    chart = ScatterChart()
                else:
                    chart = BarChart()

                chart.title = chart_def.get("title", "")
                chart.width = chart_def.get("width", 15)
                chart.height = chart_def.get("height", 10)

                if chart_def.get("style"):
                    chart.style = chart_def["style"]

                if chart_def.get("x_axis_title"):
                    chart.x_axis.title = chart_def["x_axis_title"]
                if chart_def.get("y_axis_title"):
                    chart.y_axis.title = chart_def["y_axis_title"]

                if dr:
                    data_ref = Reference(ws,
                        min_col=dr.get("min_col", 1), max_col=dr.get("max_col", 1),
                        min_row=dr.get("min_row", 1), max_row=dr.get("max_row", 1)
                    )
                    if chart_type == "pie":
                        chart.add_data(data_ref, titles_from_data=True)
                    else:
                        chart.add_data(data_ref, titles_from_data=True)

                if cr and chart_type != "scatter":
                    cats = Reference(ws,
                        min_col=cr.get("min_col", 1),
                        min_row=cr.get("min_row", 2),
                        max_row=cr.get("max_row", 10)
                    )
                    chart.set_categories(cats)

                # Show data labels for pie charts
                if chart_type == "pie":
                    chart.dataLabels = DataLabelList()
                    chart.dataLabels.showPercent = True

                ws.add_chart(chart, chart_def.get("position", "E2"))

            # --- Data Validations ---
            for dv_def in sheet_info.get("data_validations", []):
                dv_range = dv_def.get("range", "A1:A1")
                dv_type = dv_def.get("type", "list")

                if dv_type == "list":
                    formula = dv_def.get("formula", "")
                    dv = DataValidation(type="list", formula1=f'"{formula}"', allow_blank=True)
                    dv.error = dv_def.get("error_message", "Invalid input")
                    dv.errorTitle = "Validation Error"
                    ws.add_data_validation(dv)
                    dv.add(dv_range)

                elif dv_type in ("whole", "decimal"):
                    dv = DataValidation(
                        type=dv_type,
                        operator="between",
                        formula1=str(dv_def.get("min", 0)),
                        formula2=str(dv_def.get("max", 100)),
                        allow_blank=True
                    )
                    dv.error = dv_def.get("error_message", f"Value must be between {dv_def.get('min', 0)} and {dv_def.get('max', 100)}")
                    ws.add_data_validation(dv)
                    dv.add(dv_range)

            # --- Freeze panes ---
            if sheet_info.get("freeze_panes"):
                ws.freeze_panes = sheet_info["freeze_panes"]

            # --- Auto filter ---
            if sheet_info.get("auto_filter"):
                ws.auto_filter.ref = sheet_info["auto_filter"]

            # --- Print settings ---
            ps = sheet_info.get("print_settings", {})
            if ps.get("orientation") == "landscape":
                ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
            if ps.get("fit_to_page"):
                ws.page_setup.fitToPage = True
                ws.page_setup.fitToWidth = 1
                ws.page_setup.fitToHeight = 0

        wb.save(full_path)
        return {"success": True, "file_path": full_path, "message": f"Excel document saved to {file_path}"}
    except Exception as e:
        return {"error": f"Excel generation error: {str(e)}"}
